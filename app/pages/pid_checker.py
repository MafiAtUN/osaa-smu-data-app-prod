"""PID (Project Initiation Document) checker page."""

import streamlit as st
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate

from app.components.styles import page_header, section_divider
from app.services.llm_service import get_llm, get_selected_llm_name

try:
    from pypdf import PdfReader
except ImportError:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import textract
except ImportError:
    textract = None


def _extract_pdf(f):
    if PdfReader is None:
        raise ImportError("pypdf or PyPDF2 is required. Install with: pip install pypdf")
    return "".join(p.extract_text() for p in PdfReader(f).pages)


def _extract_docx(f):
    if Document is None:
        raise ImportError("python-docx is required. Install with: pip install python-docx")
    doc = Document(f)
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _extract_doc(f):
    if textract is None:
        raise ImportError("textract is required. Install with: pip install textract")
    return textract.process(f).decode("utf-8")


_CRITERIA = """
### 1. Coherence with Thematic Guidance and Priorities:
**1.1 Alignment with Priorities**: Does the PID clearly align with the USG's thematic guidance?
**1.2 Relevance to Clusters**: Does the PID address the needs of the relevant thematic cluster(s)?
**1.3 Consistency**: Are activities, outputs, outcomes, and objectives consistent with strategic guidance?

### 2. Application of Value Chain:
**2.1 Value Chain Integration**: Does the PID integrate a value chain approach?
**2.2 Outcome Focus**: Are outcomes clearly defined and achievable?
**2.3 Activity-Outcome Linkage**: Do activities link to expected outputs and outcomes?

### 3. Intentionality and Target Audience:
**3.1 Intentionality**: Is there clear intention behind each deliverable?
**3.2 Target Audience**: Does the PID identify and address the target audience?

### 4. Cross-Functional Collaboration:
**4.1 Collaborative Efforts**: Are strategies for cross-functional collaboration outlined?
**4.2 Coordination**: Does the PID facilitate coordination between teams?

### 5. Delivery of Activities:
**5.1 Role Clarity**: Are roles and responsibilities defined?
**5.2 Activity Planning**: Are activities well-structured and feasible?
**5.3 Costing**: Are costs included for each activity?
**5.4 Timeline**: Does the PID include a realistic timeline?

### 6. Monitoring and Evaluation:
**6.1 Performance Indicators**: Are there clear indicators to measure success?

### 7. Gender Mainstreaming, Disability Inclusion, Multilingualism:
**7.1 Inclusion**: Does the PID address gender, disability, and multilingualism?
"""

_THEMATIC = {
    1: "### Financing for Development\nAccess to financing, combating illicit financial flows, enhancing tax cooperation, engaging credit rating agencies, and reducing remittance costs.",
    2: "### Sustainable Development to Promote Sustainable Peace\nInclusive and equal institutional practices by countering conflict economies and building cohesive diverse societies.",
    3: "### Democracy Resilience and Human Capital\nHuman capital at the center of policy-making in Africa toward resilient societies.",
    4: "### Science Technology and Innovation\nClosing the digital divide; tackling intellectual property rights for leapfrog development.",
    5: "### Industrialization, Demographic Dividend, and AfCFTA\nHarness the demographic dividend to stimulate industrialization through the African Continental Free Trade Area.",
    6: "### Energy and Climate Action\nAfrica's energy mix, climate change, green growth, energy transition, and environmental policies.",
}

_GUIDANCE = {
    1: "Shift the paradigm of Finance for Development by tackling the finance paradox and making domestic resource mobilization the game changer.",
    2: "Durable peace requires sustainable development — not the other way around. Address root causes that prevent Africa from delivering the SDGs.",
    3: "Governance, Human capital, and Innovation: the enablers and key success factors for Africa's sustainable development.",
    4: "Governance, Human capital, and Innovation: enablers for leveraging STI for industrialization and economic transformation.",
    5: "Sustainable development requires high levels of resilience: the critical role of food systems and import substitution industrialization.",
    6: "Sustainable development requires full control of the nexus: energy planning, energy mix, and energy access.",
}

_CLUSTER_LABELS = {
    1: "1 — Financing for Development",
    2: "2 — Sustainable Development & Peace",
    3: "3 — Democracy, Resilience & Human Capital",
    4: "4 — Science, Technology & Innovation",
    5: "5 — Industrialization, Demographic Dividend & AfCFTA",
    6: "6 — Energy & Climate Action",
}

page_header(
    "📋",
    "PID Checker",
    "Upload a Project Initiation Document and evaluate it against OSAA's 7-dimension criteria framework.",
)

st.caption(f"Using: **{get_selected_llm_name()}**")
llm = get_llm()

# --- Upload ---
section_divider("Upload Document")
uploaded = st.file_uploader(
    "Upload a PID",
    type=["pdf", "doc", "docx"],
    label_visibility="collapsed",
    help="Supported formats: PDF, Word (.docx, .doc)",
)
extracted_text = None
if uploaded is not None:
    if uploaded.type == "application/pdf":
        extracted_text = _extract_pdf(uploaded)
    elif uploaded.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        extracted_text = _extract_docx(uploaded)
    elif uploaded.type == "application/msword":
        extracted_text = _extract_doc(uploaded)
    else:
        st.error("Unsupported file type.")

    if extracted_text:
        st.success(f"Document loaded — {len(extracted_text):,} characters extracted.")

# --- Cluster selection ---
section_divider("Thematic Cluster")
cluster = st.selectbox(
    "Cluster:",
    options=list(_CLUSTER_LABELS.keys()),
    index=None,
    format_func=lambda x: _CLUSTER_LABELS[x],
    placeholder="Select the PID's thematic cluster...",
    label_visibility="collapsed",
)

if cluster:
    thematic = _THEMATIC.get(cluster, "")
    guidance = _GUIDANCE.get(cluster, "")
    with st.expander("📖 Cluster Context", expanded=False):
        st.markdown(thematic)
        st.markdown(f"**Strategic guidance:** {guidance}")

thematic = _THEMATIC.get(cluster, "No thematic description available.") if cluster else ""
guidance = _GUIDANCE.get(cluster, "No strategic guidance available.") if cluster else ""

prompt = ChatPromptTemplate.from_messages([
    ("system",
     "Evaluate the PID against the criteria section by section. Use the thematic "
     "description and strategic guidance to determine alignment. Prioritize the "
     "strategic guidance. Provide reasons for each section and sub-point."),
    ("human", "PID: {document}"),
    ("human", "Criteria: {criteria}"),
    ("human", "Thematic description: {thematic_description}"),
    ("human", "Strategic guidance: {strategic_guidance}"),
])

chain = prompt | llm | StrOutputParser()

# --- Run evaluation ---
section_divider("Evaluate")
if st.button("Check PID", use_container_width=True, type="primary"):
    if not extracted_text:
        st.error("Please upload a document first.")
    elif cluster is None:
        st.error("Please select a thematic cluster.")
    else:
        with st.spinner("Evaluating PID against criteria..."):
            response = chain.invoke({
                "document": extracted_text,
                "criteria": _CRITERIA,
                "thematic_description": thematic,
                "strategic_guidance": guidance,
            })
        st.markdown("#### Evaluation Results")
        st.write(response)
